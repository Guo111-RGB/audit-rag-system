# app/main.py
import sys
import os
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
import re
from fastapi import FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel, Field
from typing import List, Optional
import uvicorn
from fastapi import UploadFile, File
import io
# ========== 1. 加载环境变量（读取 .env 中的 API Key）==========
from dotenv import load_dotenv
from app.agent.graph import build_agent_graph
from langchain_community.chat_models import ChatTongyi
from langchain_core.messages import ToolMessage
load_dotenv()  # 这会读取项目根目录下的 .env 文件

# ========== 2. 导入混合检索器 ==========
# 确保 app/hybrid_retriever.py 存在且路径正确
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from app.hybrid_retriever import HybridRetriever

# ========== 3. 配置路径 ==========
BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
CHROMA_PATH = os.path.join(BASE_DIR, "chroma_db_clause")
BM25_PATH = os.path.join(BASE_DIR, "bm25_clause.pkl")

# ========== 4. 初始化检索器 ==========
hybrid_retriever = None
HYBRID_AVAILABLE = False
if os.path.exists(CHROMA_PATH) and os.path.exists(BM25_PATH):
    try:
        hybrid_retriever = HybridRetriever(CHROMA_PATH, BM25_PATH)
        HYBRID_AVAILABLE = True
        print("✅ 混合检索器加载成功")
    except Exception as e:
        print(f"❌ 混合检索器加载失败: {e}")
else:
    print("❌ 索引文件不存在，请先运行构建脚本")

# ========== 5. 初始化通义千问大模型 ==========
LLM_AVAILABLE = False
llm = ChatTongyi(model="qwen-plus", temperature=0.3)
try:
    from langchain_community.chat_models import ChatTongyi
    from langchain_core.messages import HumanMessage, SystemMessage

    api_key = os.getenv("DASHSCOPE_API_KEY")
    if api_key and api_key.startswith("sk-"):
        llm = ChatTongyi(model="qwen-plus", temperature=0.3)
        LLM_AVAILABLE = True
        print("✅ 通义千问大模型加载成功")
    else:
        print("❌ 未找到有效的 DASHSCOPE_API_KEY，请检查 .env 文件")
except ImportError as e:
    print(f"❌ 大模型依赖未安装: {e}")
    print("   请运行: pip install langchain-community dashscope")
except Exception as e:
    print(f"❌ 大模型初始化失败: {e}")

# ========== 6. 创建 FastAPI 应用 ==========
app = FastAPI(
    title="智能审计RAG系统",
    description="基于条款分割的混合检索 + 通义千问生成",
    version="3.0",
    swagger_ui_parameters={"defaultLocale": "zh-CN"}
)
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


# ========== 7. 定义请求/响应模型 ==========
class SearchRequest(BaseModel):
    query: str
    top_k: int = 5


class SearchResult(BaseModel):
    content: str
    metadata: dict
    score: Optional[float] = None


class SearchResponse(BaseModel):
    results: List[SearchResult]


class ChatRequest(BaseModel):
    question: Optional[str] = Field(None, description="用户提出的审计问题（选填）")
    contract_clause: Optional[str] = Field(None, description="待审查的合同条款原文（选填）")
    use_llm: bool = Field(True, description="是否使用大模型生成自然语言回答")


class ChatResponse(BaseModel):
    answer: str
    references: List[str] = []


# ========== 8. API 接口 ==========
@app.get("/")
def root():
    return {"message": "智能审计RAG系统已启动", "docs": "/docs"}


@app.post("/search_full", response_model=SearchResponse)
def search_full(req: SearchRequest):
    """纯检索接口：返回最相关的条款原文"""
    if not HYBRID_AVAILABLE:
        raise HTTPException(503, detail="检索器未就绪")
    results = hybrid_retriever.hybrid_search(req.query, k=req.top_k, alpha=0.7)
    formatted = []
    for doc, score in results:
        formatted.append(SearchResult(
            content=doc.page_content,
            metadata=doc.metadata,
            score=float(score)
        ))
    return SearchResponse(results=formatted)


@app.post("/chat_full", response_model=ChatResponse)
def chat_full(req: ChatRequest):
    """
    对话接口（检索 + 大模型生成）：
    - 如果 use_llm=True 且大模型可用，生成自然语言审计报告
    - 否则只返回检索到的原文
    """
    if not HYBRID_AVAILABLE:
        raise HTTPException(503, detail="检索器未就绪，请检查索引文件")

    # --- 1. 执行混合检索 ---
    results = hybrid_retriever.hybrid_search(req.question, k=5, alpha=0.7)
    if not results:
        return ChatResponse(answer="未找到相关信息，请尝试其他问题。", references=[])

    # --- 2. 提取上下文和引用 ---
    context_parts = []
    references = []
    for doc, score in results[:4]:  # 取前4条作为上下文
        clause_id = doc.metadata.get("clause_id", "未知条款")
        source = doc.metadata.get("source", "未知来源")
        context_parts.append(f"【来源：{source} | 条款：{clause_id}】\n{doc.page_content}")
        references.append(f"{source} - {clause_id}")
    legal_context = "\n\n".join(context_parts)

    # --- 3. 判断是否使用大模型 ---
    if req.use_llm and LLM_AVAILABLE and llm:
        try:
            # 构建专业审计提示词
            system_prompt = SystemMessage(content=(
                "你是一名金融审计专家。\n"
                "请根据提供的【合同条款】和【法律法规原文】，输出审计报告。\n"
                "格式如下（严格遵守）：\n"
                "=== 审计结论 ===\n"
                "合规/部分合规/不合规\n"
                "\n"
                "=== 风险清单 ===\n"
                "1. 风险点：xxx → 违反【法条编号】\n"
                "2. 风险点：xxx → 违反【法条编号】\n"
                "（最多列出5条，每条一行，不展开论述）\n"
                "\n"
                "=== 修改建议 ===\n"
                "1. 建议xxx\n"
                "2. 建议xxx\n"
                "（每条一句话，不超过50字）\n"
                "\n"
                "【重要】只输出以上格式，不要添加任何额外内容、解释、说明或开场白。"
            ))
            human_prompt = HumanMessage(content=(
                f"用户问题：{req.question}\n\n"
                f"检索到的相关法律条款原文：\n{legal_context}"
            ))

            response = llm.invoke([system_prompt, human_prompt])
            answer = response.content

            # 如果回答太短或明显是错误信息，降级处理
            if len(answer) < 10:
                answer = f"大模型生成内容过短，以下是检索到的原文：\n{legal_context}"

        except Exception as e:
            # 大模型调用失败（如网络超时、余额不足），降级为纯检索
            answer = f"⚠️ 大模型生成失败（错误：{str(e)}），以下是检索到的相关条款原文：\n\n{legal_context}"
    else:
        # 不使用大模型 或 大模型不可用
        if not LLM_AVAILABLE:
            reason = "大模型未配置（请检查 .env 中的 DASHSCOPE_API_KEY）"
        else:
            reason = "用户设置了 use_llm=false"
        answer = f"【纯检索模式 - {reason}】\n\n检索到的相关条款原文：\n{legal_context}"

    return ChatResponse(answer=answer, references=references)


@app.get("/health")
def health():
    return {
        "status": "ok",
        "hybrid_available": HYBRID_AVAILABLE,
        "llm_available": LLM_AVAILABLE
    }


@app.post("/audit_contract", response_model=ChatResponse)
def audit_contract(req: ChatRequest):

    """
    专用合同审计接口：
    1. 接收合同条款
    2. 从法律知识库检索相关法条
    3. 大模型进行对比分析，输出风险报告
    """
    if not HYBRID_AVAILABLE:
        raise HTTPException(503, detail="检索器未就绪，请检查索引文件")

    # 如果没有提供合同条款，给出提示
    if not req.contract_clause:
        return ChatResponse(
            answer="请提供待审查的合同条款（contract_clause）",
            references=[]
        )

    # --- 1. 用合同条款作为查询词，检索相关法律 ---
    # 注意：这里直接用合同条款内容去检索，找到最相关的法律条文
    results = hybrid_retriever.hybrid_search(req.contract_clause, k=5, alpha=0.7)
    if not results:
        return ChatResponse(
            answer="未找到与您合同相关的法律法规，请检查合同条款是否涉及金融或民事法律范畴。",
            references=[]
        )

    # --- 2. 构造上下文 ---
    context_parts = []
    references = []
    for doc, score in results[:4]:
        clause_id = doc.metadata.get("clause_id", "未知条款")
        source = doc.metadata.get("source", "未知来源")
        context_parts.append(f"【来源：{source} | 条款：{clause_id}】\n{doc.page_content}")
        references.append(f"{source} - {clause_id}")
    legal_context = "\n\n".join(context_parts)

    # --- 3. 如果大模型不可用，直接返回检索到的法条 ---
    if not req.use_llm or not LLM_AVAILABLE or not llm:
        return ChatResponse(
            answer=f"【纯检索模式】\n\n您提供的合同条款：\n{req.contract_clause}\n\n相关法律条文参考：\n{legal_context}",
            references=references
        )

    # --- 4. 调用大模型进行专业审计对比 ---
    try:
        # 构造专门的审计提示词
        system_prompt = SystemMessage(content=(
            "你是一名资深金融审计与法律合规专家。\n"
            "【任务】审查用户提供的【合同条款】，与你检索到的【法律法规原文】进行逐条对比。\n"
            "【输出要求】请严格按照以下格式输出审计报告：\n"
            "1. 【合规性评估】：指出合同条款中哪些内容符合法律要求。\n"
            "2. 【风险点识别】：逐条列出合同中存在的法律风险或不合规之处，并标明对应的法律依据（引用具体法条编号）。\n"
            "3. 【修改建议】：针对每个风险点，给出具体的修改建议或合规表述。\n"
            "4. 【总体结论】：给出该条款整体的合规结论（合规/部分合规/不合规）。\n"
            "【重要原则】必须基于提供的法律原文，严禁编造法律依据。如果法律原文未涉及相关内容，请如实说明。"
        ))

        human_prompt = HumanMessage(content=(
            f"【待审查的合同条款】\n{req.contract_clause}\n\n"
            f"【参考法律法规原文】\n{legal_context}"
        ))

        response = llm.invoke([system_prompt, human_prompt])
        answer = response.content

        # 如果回答太短，降级
        if len(answer) < 20:
            answer = f"大模型生成内容过短，以下是检索到的相关法规原文供您参考：\n{legal_context}"

    except Exception as e:
        answer = f"⚠️ 大模型审计生成失败（错误：{str(e)}），以下是相关法规原文供您参考：\n{legal_context}"

    return ChatResponse(answer=answer, references=references)


@app.post("/audit_contract_file")
async def audit_contract_file(
        file: UploadFile = File(...),
        use_llm: bool = True
):
    """上传合同文档（PDF/Word/TXT），自动提取文本并生成审计报告"""

    import io
    from fastapi import UploadFile, File

    content = await file.read()
    text = ""
    filename = file.filename.lower()

    try:
        if filename.endswith(".pdf"):
            from pypdf import PdfReader
            reader = PdfReader(io.BytesIO(content))
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        elif filename.endswith(".docx"):
            from docx import Document
            doc = Document(io.BytesIO(content))
            for para in doc.paragraphs:
                text += para.text + "\n"
        elif filename.endswith(".txt"):
            text = content.decode("utf-8")
        else:
            return {"error": f"不支持的文件类型: {filename}"}
    except Exception as e:
        return {"error": f"文件解析失败: {str(e)}"}

    if not text.strip():
        return {"error": "未能提取到文本内容"}

    MAX_LENGTH = 5000
    truncated_text = text[:MAX_LENGTH]
    if len(text) > MAX_LENGTH:
        truncated_text += f"\n\n...（原文档共 {len(text)} 字符，仅取前 {MAX_LENGTH} 字符）"

    # 调用原有的审计接口
    return audit_contract(
        ChatRequest(
            contract_clause=truncated_text,
            use_llm=use_llm
        )
    )


class AgentRequest(BaseModel):
    query: str


@app.post("/agent/audit")
async def agent_audit(request: AgentRequest):
    try:
        query = request.query
        if not query:
            return {"error": "请提供审计问题或合同条款"}

        agent_graph = build_agent_graph(llm)
        initial_state = {
            "messages": [HumanMessage(content=query)],
            "next_action": "",
            "tool_results": []
        }
        config = {"recursion_limit": 10}
        result = agent_graph.invoke(initial_state, config)

        final_messages = result.get("messages", [])
        final_answer = final_messages[-1].content if final_messages else "未生成有效回答"

        references = []
        for msg in final_messages:
            if isinstance(msg, ToolMessage):
                references.append(msg.content[:200])

        return {"answer": final_answer, "references": references}
    except Exception as e:
        import traceback
        error_trace = traceback.format_exc()
        print(error_trace)  # 终端打印
        return {
            "error": str(e),
            "traceback": error_trace
        }, 500


# 这段代码需要放在 app/main.py 的合适位置
# 确保文件顶部已经导入了以下内容（如果已存在则无需重复导入）：
# from fastapi import UploadFile, File
# import io

@app.post("/agent/audit_file")
async def agent_audit_file(
        file: UploadFile = File(...),
        use_llm: bool = True
):
    """
    Agent 智能审计接口 - 支持文件上传（PDF/Word/TXT）
    """
    # 1. 读取并提取文件内容
    content = await file.read()
    text = ""
    filename = file.filename.lower()

    try:
        if filename.endswith(".pdf"):
            from pypdf import PdfReader
            reader = PdfReader(io.BytesIO(content))
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        elif filename.endswith(".docx"):
            from docx import Document
            doc = Document(io.BytesIO(content))
            for para in doc.paragraphs:
                text += para.text + "\n"
        elif filename.endswith(".txt"):
            text = content.decode("utf-8")
        else:
            return {"error": f"不支持的文件类型: {filename}"}
    except Exception as e:
        return {"error": f"文件解析失败: {str(e)}"}

    if not text.strip():
        return {"error": "未能提取到文本内容"}

    # 2. 限制长度（防止超出上下文）
    MAX_LENGTH = 5000
    truncated_text = text[:MAX_LENGTH]
    if len(text) > MAX_LENGTH:
        truncated_text += f"\n\n...（原文档共 {len(text)} 字符，仅取前 {MAX_LENGTH} 字符）"

    # 3. 调用 Agent 审计
    agent_graph = build_agent_graph(llm)
    initial_state = {
        "messages": [HumanMessage(content=truncated_text)],
        "next_action": "",
        "tool_results": []
    }
    config = {"recursion_limit": 20}
    result = agent_graph.invoke(initial_state, config)

    final_messages = result.get("messages", [])
    final_answer = final_messages[-1].content if final_messages else "未生成有效回答"

    return {"answer": final_answer}



if __name__ == "__main__":
    uvicorn.run(app, host="0.0.0.0", port=8000, log_level="debug")